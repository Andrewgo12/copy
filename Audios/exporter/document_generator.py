import os
import asyncio
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Handles document generation in multiple formats"""
    
    def __init__(self):
        self.output_dir = Path(os.getenv("UPLOAD_FOLDER", "uploads")) / "exports"
        self.output_dir.mkdir(exist_ok=True)
        
        self.formats = {
            "APA": self._format_apa,
            "BOEM": self._format_boem,
            "FREE": self._format_free
        }
    
    async def generate_document(
        self, 
        transcription_data: Dict[str, Any], 
        format_type: str, 
        style: str, 
        task_id: str
    ) -> Path:
        """
        Generate document in specified format
        
        Args:
            transcription_data: Transcription result data
            format_type: pdf, docx, or txt
            style: APA, BOEM, or FREE
            task_id: Task identifier
            
        Returns:
            Path to generated document
        """
        try:
            # Format the content according to style
            formatted_content = await self._format_content(transcription_data, style)
            
            # Generate document based on format type
            if format_type.lower() == "pdf":
                return await self._generate_pdf(formatted_content, task_id, style)
            elif format_type.lower() == "docx":
                return await self._generate_docx(formatted_content, task_id, style)
            elif format_type.lower() == "txt":
                return await self._generate_txt(formatted_content, task_id, style)
            else:
                raise ValueError(f"Unsupported format: {format_type}")
                
        except Exception as e:
            logger.error(f"Error generating document: {e}")
            raise
    
    async def _format_content(self, data: Dict[str, Any], style: str) -> Dict[str, Any]:
        """Format content according to specified style"""
        try:
            formatter = self.formats.get(style, self._format_free)
            return await asyncio.get_event_loop().run_in_executor(
                None, formatter, data
            )
        except Exception as e:
            logger.error(f"Error formatting content: {e}")
            raise
    
    def _format_apa(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Format content in APA style"""
        try:
            speakers = data.get("speakers", [])
            duration = data.get("duration", 0)
            
            # APA style header
            title = "Transcripción de Audio"
            subtitle = f"Duración: {self._format_duration(duration)}"
            
            # Format speakers in APA style
            formatted_speakers = []
            
            for speaker in speakers:
                speaker_name = speaker.get("name", "Participante")
                segments = speaker.get("segments", [])
                
                formatted_segments = []
                for segment in segments:
                    # APA format: [Timestamp] Text.
                    timestamp = segment.get("timestamp", "00:00:00")
                    text = segment.get("text", "").strip()
                    
                    if text:
                        # Ensure proper punctuation
                        if not text.endswith(('.', '!', '?')):
                            text += '.'
                        
                        formatted_segments.append({
                            "timestamp": timestamp,
                            "text": text,
                            "confidence": segment.get("confidence", 0)
                        })
                
                formatted_speakers.append({
                    "name": speaker_name,
                    "segments": formatted_segments,
                    "total_duration": speaker.get("total_duration", 0),
                    "word_count": speaker.get("word_count", 0)
                })
            
            return {
                "title": title,
                "subtitle": subtitle,
                "style": "APA",
                "speakers": formatted_speakers,
                "metadata": {
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "total_duration": duration,
                    "total_speakers": len(speakers),
                    "format_description": "Formato APA - American Psychological Association"
                }
            }
            
        except Exception as e:
            logger.error(f"Error in APA formatting: {e}")
            raise
    
    def _format_boem(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Format content in BOEM (interview) style"""
        try:
            speakers = data.get("speakers", [])
            duration = data.get("duration", 0)
            
            # BOEM style header
            title = "Transcripción de Entrevista"
            subtitle = f"Duración total: {self._format_duration(duration)}"
            
            # Format speakers in BOEM style (interview format)
            formatted_speakers = []
            
            for i, speaker in enumerate(speakers):
                # Use more formal naming for BOEM
                if i == 0:
                    speaker_name = "Entrevistador"
                else:
                    speaker_name = f"Entrevistado {i}"
                
                segments = speaker.get("segments", [])
                formatted_segments = []
                
                for segment in segments:
                    timestamp = segment.get("timestamp", "00:00:00")
                    text = segment.get("text", "").strip()
                    
                    if text:
                        # BOEM format: more conversational
                        formatted_segments.append({
                            "timestamp": timestamp,
                            "text": text,
                            "confidence": segment.get("confidence", 0)
                        })
                
                formatted_speakers.append({
                    "name": speaker_name,
                    "segments": formatted_segments,
                    "total_duration": speaker.get("total_duration", 0),
                    "word_count": speaker.get("word_count", 0)
                })
            
            return {
                "title": title,
                "subtitle": subtitle,
                "style": "BOEM",
                "speakers": formatted_speakers,
                "metadata": {
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "total_duration": duration,
                    "total_speakers": len(speakers),
                    "format_description": "Formato BOEM - Entrevistas y conversaciones"
                }
            }
            
        except Exception as e:
            logger.error(f"Error in BOEM formatting: {e}")
            raise
    
    def _format_free(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Format content in free style"""
        try:
            speakers = data.get("speakers", [])
            duration = data.get("duration", 0)
            
            # Free style header
            title = "Transcripción de Audio"
            subtitle = f"Procesado el {datetime.now().strftime('%d/%m/%Y')} - Duración: {self._format_duration(duration)}"
            
            # Keep original speaker names and format
            formatted_speakers = []
            
            for speaker in speakers:
                speaker_name = speaker.get("name", "Hablante")
                segments = speaker.get("segments", [])
                
                formatted_segments = []
                for segment in segments:
                    timestamp = segment.get("timestamp", "00:00:00")
                    text = segment.get("text", "").strip()
                    
                    if text:
                        formatted_segments.append({
                            "timestamp": timestamp,
                            "text": text,
                            "confidence": segment.get("confidence", 0)
                        })
                
                formatted_speakers.append({
                    "name": speaker_name,
                    "segments": formatted_segments,
                    "total_duration": speaker.get("total_duration", 0),
                    "word_count": speaker.get("word_count", 0)
                })
            
            return {
                "title": title,
                "subtitle": subtitle,
                "style": "FREE",
                "speakers": formatted_speakers,
                "metadata": {
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "total_duration": duration,
                    "total_speakers": len(speakers),
                    "format_description": "Formato libre - Personalizable"
                }
            }
            
        except Exception as e:
            logger.error(f"Error in free formatting: {e}")
            raise
    
    async def _generate_pdf(self, content: Dict[str, Any], task_id: str, style: str) -> Path:
        """Generate PDF document"""
        try:
            output_path = self.output_dir / f"transcription_{task_id}_{style}.pdf"
            
            # Create PDF document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center
            )
            
            speaker_style = ParagraphStyle(
                'SpeakerName',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                textColor=colors.darkblue
            )
            
            timestamp_style = ParagraphStyle(
                'Timestamp',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.grey
            )
            
            # Build document content
            story = []
            
            # Title and subtitle
            story.append(Paragraph(content["title"], title_style))
            story.append(Paragraph(content["subtitle"], styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Metadata table
            metadata = content["metadata"]
            meta_data = [
                ["Generado:", metadata["generated_at"]],
                ["Duración total:", self._format_duration(metadata["total_duration"])],
                ["Número de voces:", str(metadata["total_speakers"])],
                ["Formato:", metadata["format_description"]]
            ]
            
            meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
            meta_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(meta_table)
            story.append(Spacer(1, 30))
            
            # Speakers and segments
            for speaker in content["speakers"]:
                # Speaker name
                story.append(Paragraph(speaker["name"], speaker_style))
                
                # Speaker segments
                for segment in speaker["segments"]:
                    # Timestamp
                    story.append(Paragraph(f"[{segment['timestamp']}]", timestamp_style))
                    
                    # Text
                    story.append(Paragraph(segment["text"], styles['Normal']))
                    story.append(Spacer(1, 12))
                
                story.append(Spacer(1, 20))
            
            # Build PDF
            await asyncio.get_event_loop().run_in_executor(
                None, doc.build, story
            )
            
            logger.info(f"PDF generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise
    
    async def _generate_docx(self, content: Dict[str, Any], task_id: str, style: str) -> Path:
        """Generate DOCX document"""
        try:
            output_path = self.output_dir / f"transcription_{task_id}_{style}.docx"
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading(content["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            subtitle = doc.add_paragraph(content["subtitle"])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_heading("Información del documento", level=2)
            metadata = content["metadata"]
            
            meta_table = doc.add_table(rows=4, cols=2)
            meta_table.style = 'Table Grid'
            
            meta_data = [
                ("Generado:", metadata["generated_at"]),
                ("Duración total:", self._format_duration(metadata["total_duration"])),
                ("Número de voces:", str(metadata["total_speakers"])),
                ("Formato:", metadata["format_description"])
            ]
            
            for i, (key, value) in enumerate(meta_data):
                meta_table.cell(i, 0).text = key
                meta_table.cell(i, 1).text = value
            
            doc.add_paragraph()
            
            # Transcription content
            doc.add_heading("Transcripción", level=2)
            
            for speaker in content["speakers"]:
                # Speaker heading
                speaker_heading = doc.add_heading(speaker["name"], level=3)
                
                # Speaker segments
                for segment in speaker["segments"]:
                    # Timestamp paragraph
                    timestamp_para = doc.add_paragraph()
                    timestamp_run = timestamp_para.add_run(f"[{segment['timestamp']}]")
                    timestamp_run.font.size = Pt(9)
                    # Set gray color for timestamp
                    from docx.shared import RGBColor
                    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
                    
                    # Text paragraph
                    text_para = doc.add_paragraph(segment["text"])
                    text_para.style = 'Normal'
                
                doc.add_paragraph()
            
            # Save document
            await asyncio.get_event_loop().run_in_executor(
                None, doc.save, str(output_path)
            )
            
            logger.info(f"DOCX generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise
    
    async def _generate_txt(self, content: Dict[str, Any], task_id: str, style: str) -> Path:
        """Generate TXT document"""
        try:
            output_path = self.output_dir / f"transcription_{task_id}_{style}.txt"
            
            # Build text content
            text_content = []
            
            # Header
            text_content.append("=" * 60)
            text_content.append(content["title"].upper())
            text_content.append("=" * 60)
            text_content.append("")
            text_content.append(content["subtitle"])
            text_content.append("")
            
            # Metadata
            metadata = content["metadata"]
            text_content.append("INFORMACIÓN DEL DOCUMENTO")
            text_content.append("-" * 30)
            text_content.append(f"Generado: {metadata['generated_at']}")
            text_content.append(f"Duración total: {self._format_duration(metadata['total_duration'])}")
            text_content.append(f"Número de voces: {metadata['total_speakers']}")
            text_content.append(f"Formato: {metadata['format_description']}")
            text_content.append("")
            
            # Transcription
            text_content.append("TRANSCRIPCIÓN")
            text_content.append("-" * 30)
            text_content.append("")
            
            for speaker in content["speakers"]:
                # Speaker name
                text_content.append(f">>> {speaker['name'].upper()} <<<")
                text_content.append("")
                
                # Segments
                for segment in speaker["segments"]:
                    text_content.append(f"[{segment['timestamp']}]")
                    text_content.append(segment["text"])
                    text_content.append("")
                
                text_content.append("-" * 40)
                text_content.append("")
            
            # Write to file
            final_content = "\n".join(text_content)

            # Write file in executor to avoid blocking
            await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: output_path.write_text(final_content, encoding='utf-8')
            )
            
            logger.info(f"TXT generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating TXT: {e}")
            raise
    
    def _format_duration(self, seconds: float) -> str:
        """Format duration in HH:MM:SS format"""
        try:
            hours = int(seconds // 3600)
            minutes = int((seconds % 3600) // 60)
            secs = int(seconds % 60)
            
            if hours > 0:
                return f"{hours:02d}:{minutes:02d}:{secs:02d}"
            else:
                return f"{minutes:02d}:{secs:02d}"
        except:
            return "00:00"
    
    async def cleanup_exports(self, task_id: str):
        """Clean up exported files for a task"""
        try:
            pattern = f"transcription_{task_id}_*"
            for file_path in self.output_dir.glob(pattern):
                file_path.unlink()
            logger.info(f"Cleaned up exports for task {task_id}")
        except Exception as e:
            logger.warning(f"Error cleaning up exports: {e}")
